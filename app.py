
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import tempfile
import os
import datetime
import re

st.set_page_config(page_title="Gerador de Propostas", layout="centered")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
kits_file = os.path.join(BASE_DIR, 'kits.xlsx')
modelo_default = os.path.join(BASE_DIR, 'modelo_novo.docx')
output_dir = os.path.join(BASE_DIR, 'propostas_geradas')
logo_path = os.path.join(BASE_DIR, 'imagens', 'logo.png')

def formatar_moeda(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def aplicar_negrito(paragrafo, substituicoes):
    texto_original = paragrafo.text
    cor_original = None
    if paragrafo.runs:
        cor_original = paragrafo.runs[0].font.color.rgb

    for chave, valor in substituicoes.items():
        texto_original = texto_original.replace(chave, f"§§§{valor}§§§")

    if texto_original == paragrafo.text:
        return

    partes = re.split(r'(§§§.*?§§§|R\$ [\d\.,]+)', texto_original)
    paragrafo.clear()
    for parte in partes:
        run = paragrafo.add_run(parte.replace("§§§", ""))
        if parte.startswith("§§§") and parte.endswith("§§§"):
            run.bold = True
        elif re.match(r'R\$ [\d\.,]+', parte):
            run.bold = True
        if cor_original:
            run.font.color.rgb = cor_original

def gerar_proposta(kits_file, modelo_file, kits_selecionados, nome_cliente, distancia_loja, desconto_percentual, planta_files):
    df = pd.read_excel(kits_file)

    peso_total = 0
    preco_normal = 0
    valor_avista = 0
    area_num = 0
    peso_kit = 0
    descricao_kit = ""
    cod_kit = ""
    link_kit = ""
    quant = 1

    for item in kits_selecionados:
        linha = df[df['DESCRICAO'] == item['kit']]
        if linha.empty:
            continue
        peso_kit = float(linha.iloc[0]['PESO UND'])
        preco_normal = float(linha.iloc[0]['A VISTA'])
        descricao_kit = linha.iloc[0]['DESCRICAO']
        cod_kit = str(linha.iloc[0]['CODIGO'])
        link_kit = linha.iloc[0]['LINK_KIT']
        quant = item['quantidade']
        try:
            area_num = float(str(linha.iloc[0]['AREA']).replace(',', '.'))
        except (ValueError, TypeError):
            area_num = 0.0
        peso_total += peso_kit * quant
        valor_avista = preco_normal * quant * (1 - desconto_percentual / 100)

    frete_normal = (peso_total / 1000) * 1150
    frete_adicional = max(0, (distancia_loja - 200)) * 5.50
    frete_total = frete_normal + frete_adicional

    valor_chave_mao = preco_normal * quant * 2.15
    mao_obra_e_opcionais = valor_chave_mao - valor_avista

    CUB_ALVENARIA = 2900
    CUB_PREFAB = 2150

    custo_alvenaria = CUB_ALVENARIA * area_num
    custo_chave_mao = CUB_PREFAB * area_num
    economia_cub = custo_alvenaria - custo_chave_mao

    prazo_construcao = f"{int(area_num)} dias" if area_num > 0 else "Não informado"

    substituicoes = {
        '{{cod_kit}}': cod_kit,
        '{{quant}}': str(quant),
        '{{nome_cliente}}': nome_cliente,
        '{{descrição_kit}}': descricao_kit,
        '{{preço_normal}}': formatar_moeda(preco_normal),
        '{{valor_total}}': formatar_moeda(preco_normal * quant),
        '{{valor_avista}}': formatar_moeda(valor_avista),
        '{{peso_total}}': f"{peso_total:.2f} kg",
        '{{link_kit}}': link_kit,
        '{{distancia_loja}}': f"{distancia_loja} km",
        '{{frete_normal}}': formatar_moeda(frete_normal),
        '{{frete_adicional}}': formatar_moeda(frete_adicional),
        '{{frete_total}}': formatar_moeda(frete_total),
        '{{50%_valor_avista}}': formatar_moeda(valor_avista / 2),
        '{{valor_chave_mao1}}': formatar_moeda(valor_chave_mao),
        '{{valor_chave_mao}}': formatar_moeda(valor_chave_mao),
        '{{valor_kit}}': formatar_moeda(valor_avista),  # <- corrigido aqui: valor total COM desconto
        '{{valor_mao_obra}}': formatar_moeda(mao_obra_e_opcionais),
        '{{tam_kit}}': f"{area_num} m²",
        '{{prazo_construcao}}': prazo_construcao,
        '{{planta_baixa}}': f"{cod_kit}planta",
        '{{peso_kit}}': f"{peso_kit} kg",
        '{{cub_alvenaria}}': formatar_moeda(CUB_ALVENARIA),
        '{{cub_prefab}}': formatar_moeda(CUB_PREFAB),
        '{{area_casa}}': f"{area_num}",
        '{{custo_alvenaria}}': formatar_moeda(custo_alvenaria),
        '{{custo_chave_mao}}': formatar_moeda(custo_chave_mao),
        '{{economia_cub}}': formatar_moeda(economia_cub),
        '{{porcentagem_desconto}}': f'({desconto_percentual}%)',
        '{{data_atual}}': datetime.datetime.today().strftime('%d/%m/%Y')
    }

    modelo = Document(modelo_file)

    for p in modelo.paragraphs:
        aplicar_negrito(p, substituicoes)

    for tabela in modelo.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    aplicar_negrito(par, substituicoes)

    if planta_files:
        for planta_file in planta_files[:2]:
            planta_temp_path = os.path.join(tempfile.gettempdir(), planta_file.name)
            with open(planta_temp_path, "wb") as f:
                f.write(planta_file.getbuffer())
            modelo.add_picture(planta_temp_path, width=Inches(5))

    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"Proposta_{nome_cliente.replace(' ', '_')}.docx")
    modelo.save(output_path)
    return output_path

st.title("Gerador de Propostas")

nome_cliente = st.text_input("Nome do Cliente")
distancia_loja = st.number_input("Distância da Loja (km)", min_value=0.0, step=1.0)
desconto_percentual = st.slider("Desconto (%)", min_value=1, max_value=12, value=5)
kits = []

if os.path.exists(kits_file):
    df_temp = pd.read_excel(kits_file)
    lista_modelos = df_temp['DESCRICAO'].tolist()
    quant = st.number_input("Quantidade", min_value=1, step=1)
    busca = st.text_input("Digite parte do nome do modelo:")
    resultados = [modelo for modelo in lista_modelos if busca.lower() in modelo.lower()]
    if resultados:
        modelo_selecionado = st.selectbox("Selecione o modelo encontrado:", options=resultados)
        kits.append({"kit": modelo_selecionado, "quantidade": quant})
    else:
        st.warning("Nenhum modelo encontrado. Ajuste a busca.")
else:
    st.warning(f"⚠️ Arquivo {kits_file} não encontrado.")

planta_files = st.file_uploader("Plantas Baixas", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
modelo_file = st.file_uploader("Modelo DOCX", type=["docx"])

if st.button("📄 Gerar Proposta"):
    if not nome_cliente or not kits:
        st.warning("Preencha o nome do cliente e selecione o modelo.")
    else:
        modelo = modelo_file if modelo_file else modelo_default
        try:
            caminho = gerar_proposta(kits_file, modelo, kits, nome_cliente, distancia_loja, desconto_percentual, planta_files)
            if caminho:
                st.success("✅ Proposta gerada com sucesso!")
                with open(caminho, "rb") as file:
                    st.download_button("📥 Baixar Proposta", file, file_name=os.path.basename(caminho))
        except Exception as e:
            st.error(f"Erro ao gerar proposta: {e}")
